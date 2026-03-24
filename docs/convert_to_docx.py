"""
Convert PROJECT_REPORT.md to PROJECT_REPORT.docx
"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        borders.append(element)
    tblPr.append(borders)

def create_docx():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # Read the markdown file
    md_path = os.path.join(os.path.dirname(__file__), 'PROJECT_REPORT.md')
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split into lines
    lines = content.split('\n')
    
    i = 0
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Skip HTML tags like <center>, </center>
        if line.startswith('<center>') or line.startswith('</center>') or line == '&nbsp;':
            i += 1
            continue
        
        # Skip horizontal rules
        if line == '---':
            # Add a page break for major section separators
            doc.add_page_break()
            i += 1
            continue
        
        # Handle tables
        if '|' in line and line.startswith('|'):
            # Check if this is a table separator line
            if re.match(r'^\|[\s\-:|]+\|$', line):
                i += 1
                continue
            
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            
            # Check if next line is still part of table
            next_i = i + 1
            while next_i < len(lines) and not lines[next_i].strip():
                next_i += 1
            
            if next_i >= len(lines) or not lines[next_i].strip().startswith('|'):
                # End of table, create it
                if table_rows:
                    num_cols = max(len(row) for row in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    add_table_borders(table)
                    
                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:
                                cell = table.cell(row_idx, col_idx)
                                # Clean markdown formatting
                                clean_text = cell_text.replace('**', '').replace('`', '')
                                cell.text = clean_text
                                
                                # Style the cell
                                for paragraph in cell.paragraphs:
                                    paragraph.style.font.size = Pt(10)
                                    for run in paragraph.runs:
                                        run.font.size = Pt(10)
                                        run.font.name = 'Times New Roman'
                                
                                # Header row styling
                                if row_idx == 0:
                                    set_cell_shading(cell, 'D9E2F3')
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                    
                    doc.add_paragraph()  # Add spacing after table
                
                in_table = False
                table_rows = []
            
            i += 1
            continue
        
        # Handle headings
        if line.startswith('#'):
            level = len(line.split(' ')[0])  # Count #'s
            text = line.lstrip('#').strip()
            
            # Remove markdown bold
            text = text.replace('**', '')
            
            if level == 1:
                p = doc.add_heading(text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                p = doc.add_heading(text, level=1)
                if 'Chapter' in text:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 3:
                p = doc.add_heading(text, level=2)
            elif level >= 4:
                p = doc.add_heading(text, level=3)
            
            i += 1
            continue
        
        # Handle bold paragraphs (like figure captions)
        if line.startswith('**') and line.endswith('**'):
            text = line.strip('*').strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            i += 1
            continue
        
        # Handle image references (skip - they need actual images)
        if line.startswith('!['):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Extract alt text
            alt_match = re.match(r'!\[([^\]]*)\]', line)
            if alt_match:
                run = p.add_run(f'[{alt_match.group(1)}]')
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.font.size = Pt(10)
            i += 1
            continue
        
        # Handle bullet points
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Clean markdown formatting
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # bold
            text = text.replace('`', '')
            
            p = doc.add_paragraph(style='List Bullet')
            
            # Handle bold parts within bullet
            parts = re.split(r'(\*\*[^*]+\*\*)', line[2:].strip())
            p.clear()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.strip('*'))
                    run.bold = True
                else:
                    clean = part.replace('`', '')
                    p.add_run(clean)
            
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            
            i += 1
            continue
        
        # Handle numbered items with bold labels (like "1. **Title...**")
        numbered_match = re.match(r'^(\d+)\.\s+(.+)', line)
        if numbered_match:
            text = numbered_match.group(0)
            text = text.replace('`', '')
            
            p = doc.add_paragraph(style='List Number')
            p.clear()
            
            # Handle bold parts
            remaining = text[len(numbered_match.group(1)) + 2:]  # skip "N. "
            parts = re.split(r'(\*\*[^*]+\*\*)', remaining)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.strip('*'))
                    run.bold = True
                else:
                    p.add_run(part)
            
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            
            i += 1
            continue
        
        # Regular paragraph
        # Clean markdown formatting
        clean_line = line.replace('`', '')
        
        p = doc.add_paragraph()
        
        # Handle mixed bold/normal text
        parts = re.split(r'(\*\*[^*]+\*\*)', clean_line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part.strip('*'))
                run.bold = True
            else:
                p.add_run(part)
        
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        # Check for specific alignment needs
        if any(keyword in clean_line for keyword in ['SRI VENKATESWARA', 'R.V.S NAGAR', 'DEPARTMENT OF', 'MAY 2026', 'AUTONOMOUS', 'APPROVED BY', 'ACCREDITED BY', 'AN ISO']):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        i += 1
    
    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'PROJECT_REPORT.docx')
    doc.save(output_path)
    print(f"Successfully created: {output_path}")

if __name__ == '__main__':
    create_docx()
