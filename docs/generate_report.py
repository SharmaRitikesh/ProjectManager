"""
Professional DOCX Report Generator - Matches finalproject PDF format
Generates: PROJECT_REPORT.docx with headers, footers, page borders, chapters
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import os

PROJ_TITLE = "PROJECT MANAGER – FULL STACK WEB APPLICATION USING JAVA SPRING BOOT AND REACT"
DEPT = "CSE, SVCET"

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{edge}')
        for attr, val in [('val','single'),('sz','4'),('space','0'),('color','000000')]:
            el.set(qn(f'w:{attr}'), val)
        borders.append(el)
    tblPr.append(borders)

def add_page_border(section):
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for edge in ('top','left','bottom','right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), '000000')
        pgBorders.append(border)
    sectPr.append(pgBorders)

def setup_header(section, title=PROJ_TITLE):
    header = section.header
    header.is_linked_to_previous = False
    p1 = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p1.text = title
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p1.runs:
        r.font.size = Pt(9)
        r.font.name = 'Times New Roman'
        r.bold = True
    p2 = header.add_paragraph(DEPT)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p2.runs:
        r.font.size = Pt(9)
        r.font.name = 'Times New Roman'

def setup_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.append(fld)

def add_centered(doc, text, size=14, bold=True, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    r.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_para(doc, text, size=12, bold=False, align=None, space_after=6, first_indent=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    r.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    return p

def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    return p

def add_figure(doc, image_path, caption, fig_num):
    """Insert an image with a centered caption below it."""
    diagrams_dir = os.path.join(os.path.dirname(__file__), 'diagrams')
    full_path = os.path.join(diagrams_dir, image_path)
    if os.path.exists(full_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(full_path, width=Inches(5.0))
    else:
        add_centered(doc, f'[Diagram: {caption}]', size=11, bold=False)
    # Caption
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_r = cap_p.add_run(f'Figure {fig_num}: {caption}')
    cap_r.font.size = Pt(10)
    cap_r.font.name = 'Times New Roman'
    cap_r.bold = True
    cap_p.paragraph_format.space_after = Pt(12)
    return cap_p

def add_chapter_heading(doc, chapter_num, title):
    doc.add_page_break()
    add_centered(doc, f"Chapter-{chapter_num:02d}", size=14, bold=True, space_after=6)
    add_centered(doc, title.upper(), size=14, bold=True, space_after=12)

def add_section_heading(doc, number, title):
    p = doc.add_paragraph()
    r = p.add_run(f"{number} {title}")
    r.font.size = Pt(13)
    r.font.name = 'Times New Roman'
    r.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

def add_subsection(doc, number, title):
    p = doc.add_paragraph()
    r = p.add_run(f"{number} {title}")
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    r.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

def make_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(t)
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = 'Times New Roman'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri+1, ci)
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Times New Roman'
    doc.add_paragraph()
    return t

def create_report():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        add_page_border(section)
        setup_header(section)
        setup_footer(section)

    # ===================== TITLE PAGE =====================
    for _ in range(4): doc.add_paragraph()
    add_centered(doc, PROJ_TITLE, size=16, bold=True, space_after=18)
    add_centered(doc, "A PROJECT REPORT", size=14, bold=True, space_after=12)
    add_centered(doc, "Submitted by", size=12, bold=False, space_after=8)
    add_centered(doc, "22781A05XX: RITIKESH SHARMA", size=12, bold=True, space_after=12)
    add_centered(doc, "in partial fulfilment of the award of the degree of", size=12, bold=False, space_after=4)
    add_centered(doc, "BACHELOR OF TECHNOLOGY", size=13, bold=True, space_after=4)
    add_centered(doc, "in", size=12, bold=False, space_after=4)
    add_centered(doc, "COMPUTER SCIENCE AND ENGINEERING", size=13, bold=True, space_after=12)
    add_centered(doc, "Under the Guidance of", size=12, bold=False, space_after=4)
    add_centered(doc, "Dr. Hare Ram Singh", size=12, bold=True, space_after=4)
    add_centered(doc, "Assistant Professor, CSE Department", size=12, bold=False, space_after=18)
    add_centered(doc, "SRI VENKATESWARA COLLEGE OF ENGINEERING AND TECHNOLOGY", size=12, bold=True, space_after=2)
    add_centered(doc, "(AUTONOMOUS)", size=11, bold=True, space_after=2)
    add_centered(doc, "R.V.S NAGAR, CHITTOOR – 517127 (A.P)", size=11, bold=False, space_after=2)
    add_centered(doc, "(Approved by AICTE, New Delhi, Affiliated to JNTUA, Ananthapuram)", size=10, bold=False, space_after=2)
    add_centered(doc, "(Accredited by NBA, New Delhi, NAAC 'A+', Bengaluru)", size=10, bold=False, space_after=2)
    add_centered(doc, "(An ISO 9001:2000 Certified Institution)", size=10, bold=False, space_after=8)
    add_centered(doc, "MAY 2026", size=13, bold=True)

    # ===================== CERTIFICATE =====================
    doc.add_page_break()
    add_centered(doc, "SRI VENKATESWARA COLLEGE OF ENGINEERING AND TECHNOLOGY (AUTONOMOUS)", size=12, bold=True, space_after=2)
    add_centered(doc, "R.V.S NAGAR, CHITTOOR – 517127 (A.P)", size=11, bold=False, space_after=2)
    add_centered(doc, "(APPROVED BY AICTE, NEW DELHI, AFFILIATED TO JNTUA, ANANTHAPURAM)", size=9, bold=False, space_after=2)
    add_centered(doc, "(ACCREDITED BY NBA, NEW DELHI, NAAC 'A+', BENGALURU)", size=9, bold=False, space_after=2)
    add_centered(doc, "DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING", size=12, bold=True, space_after=12)
    add_centered(doc, "CERTIFICATE", size=14, bold=True, space_after=12)
    add_para(doc, 'This is to certify that, the project entitled, "PROJECT MANAGER – FULL STACK WEB APPLICATION USING JAVA SPRING BOOT AND REACT" is a Bonafide work carried by "Ritikesh Sharma (22781A05XX)" student of Computer Science and Engineering Department in Sri Venkateswara College of Engineering and Technology (Autonomous) Chittoor in the year of 2025-2026.', first_indent=1.27)
    doc.add_paragraph()
    t = doc.add_table(rows=2, cols=2)
    t.cell(0,0).text = "SIGNATURE OF THE GUIDE"
    t.cell(0,1).text = "SIGNATURE OF THE HOD"
    t.cell(1,0).text = "Dr. Hare Ram Singh\nAssistant Professor, CSE"
    t.cell(1,1).text = "Dr. P. Jyotheeswari\nHOD, CSE\nM.Tech, Ph.D"
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'
    doc.add_paragraph()
    t2 = doc.add_table(rows=1, cols=2)
    t2.cell(0,0).text = "INTERNAL EXAMINER"
    t2.cell(0,1).text = "EXTERNAL EXAMINER"
    add_para(doc, "Viva-Voce Conducted on _____________________")

    # ===================== DECLARATION =====================
    doc.add_page_break()
    add_centered(doc, "SRI VENKATESWARA COLLEGE OF ENGINEERING AND TECHNOLOGY (AUTONOMOUS)", size=12, bold=True, space_after=2)
    add_centered(doc, "R.V.S NAGAR, CHITTOOR – 517127 (A.P)", size=11, bold=False, space_after=2)
    add_centered(doc, "DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING", size=12, bold=True, space_after=12)
    add_centered(doc, "DECLARATION", size=14, bold=True, space_after=12)
    add_para(doc, 'I, Ritikesh Sharma (22781A05XX), hereby declare that the project report entitled "PROJECT MANAGER – FULL STACK WEB APPLICATION USING JAVA SPRING BOOT AND REACT" under the guidance of Dr. Hare Ram Singh, Assistant Professor, Sri Venkateswara College of Engineering & Technology (Autonomous), Chittoor, is submitted in partial fulfillment of the requirements for the award of the degree of BACHELOR OF TECHNOLOGY in COMPUTER SCIENCE AND ENGINEERING.', first_indent=1.27)
    add_para(doc, "This is a record of Bonafide work carried out by me. The results embodied in this project report have not been reproduced or copied from any source. The results embodied in this project report have not been submitted to any other university or institute for the award of any other degree or diploma.", first_indent=1.27)
    doc.add_paragraph()
    add_para(doc, "Signature of Student", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(doc, "Ritikesh Sharma    22781A05XX", align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ===================== ACKNOWLEDGEMENT =====================
    doc.add_page_break()
    add_centered(doc, "ACKNOWLEDGEMENT", size=14, bold=True, space_after=12)
    add_para(doc, "A grateful thanks to Dr. R. VENKATASWAMY, chairman of Sri Venkateswara College of Engineering and Technology for providing education in their esteemed institution.", first_indent=1.27)
    add_para(doc, "I wish to record my deep sense of gratitude and profound thanks to our beloved Vice Chairman, Sri R. V. SRINIVAS for his valuable support throughout the course.", first_indent=1.27)
    add_para(doc, "I express my sincere thanks to Dr. M. MOHAN BABU, our beloved principal for his encouragement and suggestion during the course of study.", first_indent=1.27)
    add_para(doc, "With the deep sense of gratefulness, I acknowledge Dr. P. JYOTHEESWARI, Head of the Department, Computer Science and Engineering, for giving us her inspiring guidance in undertaking our project report.", first_indent=1.27)
    add_para(doc, "I express my sincere thanks to the Project Guide Dr. HARE RAM SINGH, Assistant Professor, Department of Computer Science and Engineering, for his keen interest, stimulating guidance, constant encouragement with our work during all stages, to bring this project into fruition.", first_indent=1.27)
    add_para(doc, "I wish to convey my gratitude and express my sincere thanks to all Project Review Committee members for their support and cooperation rendered for successful submission of our project work.", first_indent=1.27)
    add_para(doc, "Finally, I would like to express my sincere thanks to all teaching, non-teaching faculty members, my parents, friends and for all those who have supported me to complete the project work successfully.", first_indent=1.27)
    doc.add_paragraph()
    add_para(doc, "Ritikesh Sharma    22781A05XX", align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ===================== ABSTRACT =====================
    doc.add_page_break()
    add_centered(doc, "ABSTRACT", size=14, bold=True, space_after=12)
    add_para(doc, "In the modern era of software development and digital transformation, project management has become a cornerstone of successful team collaboration, workflow coordination, and timely delivery of software products. With the growing adoption of agile methodologies and distributed team structures, organizations across the globe are seeking robust, centralized, and web-based project management platforms that enable seamless task tracking, role-based access control, and real-time progress monitoring. Traditional approaches to project management, including spreadsheet-based tracking, email-driven communication, and standalone desktop tools, have proven inadequate in addressing the complexity, scalability, and accessibility demands of contemporary software development teams. These conventional methods frequently result in misaligned priorities, unclear task ownership, version conflicts, missed deadlines, and overall diminished productivity across teams.", first_indent=1.27)
    add_para(doc, "To address these significant challenges, this project presents the design, development, and implementation of a comprehensive full-stack Project Management Application built using modern, industry-standard web technologies. The backend of the application is developed using the Java Spring Boot framework, which follows a layered architecture pattern comprising Controller, Service, and Repository layers, ensuring clear separation of concerns and maintainability. Spring Data JPA is employed for database operations through an H2 relational database, managing core entities including Users, Projects, Tasks, and Project Members with well-defined relationships and cascading operations. Security is enforced through stateless JWT-based authentication integrated with Spring Security, where each API request is verified and authorized using token-based validation. The system implements a multi-tier role-based access control mechanism distinguishing between system-level roles (User, Manager, Admin) and project-level roles (Owner, Admin, Mentor, Member, Viewer), enabling fine-grained permission management across the platform.", first_indent=1.27)
    add_para(doc, "The frontend application is built using React 18 with a component-based architecture, utilizing React Context API for centralized authentication state management and Axios with request and response interceptors for seamless API communication. The user interface features a dynamic Dashboard with real-time statistics, a Projects management module with creation and status tracking, a Task Board with drag-and-drop status columns and visual progress indicators, and an Admin Dashboard for comprehensive user and project management. Through this integrated full-stack approach, the system delivers a secure, scalable, and user-friendly solution for modern team collaboration and project lifecycle management.", first_indent=1.27)

    # ===================== CONTENTS =====================
    doc.add_page_break()
    add_centered(doc, "CONTENTS", size=14, bold=True, space_after=12)
    contents = [
        ("1","INTRODUCTION","1-3"),("2","LITERATURE SURVEY","4-8"),
        ("3","EXISTING SYSTEM","9-11"),("4","PROPOSED SYSTEM","12-14"),
        ("5","SYSTEM REQUIREMENTS","15-16"),("6","METHODOLOGY","17-23"),
        ("7","MODULES USED","24-28"),("8","JAVA AND SPRING BOOT FUNDAMENTALS","29-38"),
        ("9","SYSTEM ARCHITECTURE","39-44"),("10","SYSTEM TESTING","45-52"),
        ("11","SCREENSHOTS","53-57"),("12","RESULT AND ANALYSIS","58-59"),
        ("13","CONCLUSION","60-61"),("14","FUTURE WORK","62-63"),
        ("15","REFERENCES","64-65"),
    ]
    make_table(doc, ["S.No","Title","Page No"], contents)

    # LIST OF FIGURES
    doc.add_page_break()
    add_centered(doc, "LIST OF FIGURES", size=14, bold=True, space_after=12)
    figs = [
        ("1","Data Flow Diagram – Level 0 (Context Diagram)","18"),
        ("2","Data Flow Diagram – Level 1","19"),
        ("3","System Architecture Diagram","22"),
        ("4","UML Use Case Diagram","23"),
        ("5","UML Class Diagram (Entity Relationships)","24"),
        ("6","UML Sequence Diagram (Authentication and Project Creation)","25"),
    ]
    make_table(doc, ["S.No","Figures","Page No"], figs)

    # LIST OF TABLES
    add_centered(doc, "LIST OF TABLES", size=14, bold=True, space_after=12)
    tables_list = [
        ("1","Technology Stack Selection","13"),
        ("2","Hardware Requirements","15"),
        ("3","Software Requirements","16"),
        ("4","REST API Endpoints","41"),
        ("5","System Test Cases","48"),
    ]
    make_table(doc, ["S.No","Tables","Page No"], tables_list)

    # Call content generation
    write_chapter_01(doc)
    write_chapter_02(doc)
    write_chapter_03(doc)
    write_chapter_04(doc)
    write_chapter_05(doc)
    write_chapter_06(doc)
    write_chapter_07(doc)
    write_chapter_08(doc)
    write_chapter_09(doc)
    write_chapter_10(doc)
    write_chapter_11(doc)
    write_chapter_12(doc)
    write_chapter_13(doc)
    write_chapter_14(doc)
    write_chapter_15(doc)

    out = os.path.join(os.path.dirname(__file__), 'PROJECT_REPORT_v3.docx')
    doc.save(out)
    print(f"Report saved: {out}")

# Content functions loaded from separate file
from report_content import *

if __name__ == '__main__':
    create_report()
